# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ---- 全局样式 ----
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    return p


# ==================== 封面 ====================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI 动作姿态纠错系统')
run.font.size = Pt(28)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('项目报告')
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_page_break()

# ==================== 目录占位 ====================
doc.add_heading('目录', level=1)
add_para('（请在 Word 中插入自动目录）')
doc.add_page_break()

# ==================== 1. 项目概述 ====================
doc.add_heading('1  项目概述', level=1)

add_para(
    '本项目实现了一个基于计算机视觉的动作姿态纠错系统。系统接收两段视频作为输入——'
    '一段是老师的标准示范视频，另一段是学生的练习视频——通过人体骨架提取、空间归一化、'
    '时间序列对齐和逐帧误差检测等步骤，自动生成一段左右对比的纠错视频，'
    '在其中以红色标注、箭头引导等方式直观展示学生动作与标准动作之间的偏差。', indent=True)

add_para(
    '该系统可应用于体育教学、舞蹈训练、康复训练等场景，'
    '帮助降低人工纠错成本，提供客观、可量化的动作评估反馈。', indent=True)

# ==================== 2. 系统架构 ====================
doc.add_heading('2  系统架构', level=1)

add_para('系统整体处理流程分为以下 6 个步骤：', indent=True)

steps = [
    '视频输入与预处理：读取老师和学生视频，统一缩放到 600px 高度（等比缩放）。',
    '姿态提取：使用 MediaPipe Pose 模型逐帧提取 33 个人体关键点的三维坐标。',
    '平滑降噪：对骨架坐标序列做 5 帧滑动平均滤波，消除 MediaPipe 的逐帧抖动噪声。',
    '空间归一化：以骨盆中心为原点、躯干长度为缩放因子，将所有关节转为相对坐标，消除体型差异。',
    '时间对齐（DTW）：使用动态时间规整算法对齐两个序列的帧，解决动作速度不同步的问题。',
    '误差检测与视觉渲染：逐帧比对归一化后的骨架坐标，标记错误关节，生成带纠错标注的对比视频。',
]
for i, step in enumerate(steps, 1):
    add_para(f'（{i}）{step}', indent=True)

# ==================== 3. 核心算法 ====================
doc.add_heading('3  核心算法详解', level=1)

# 3.1 姿态提取
doc.add_heading('3.1  姿态提取', level=2)

add_para(
    '系统使用 Google 开源的 MediaPipe Pose 框架进行人体姿态估计。'
    '该模型将人体建模为 33 个关节点，每个点输出 (x, y, z) 三维坐标，'
    '坐标值归一化到 0~1 范围（相对于画面尺寸）。', indent=True)

add_para(
    '33 个点涵盖脸部、躯干、上肢和下肢。在实际比对中，并非所有点都参与误差计算。'
    '脸部细节和手指等末端点噪声较大且对动作判断帮助不大，因此被排除。'
    '系统选取 12 个核心关节用于 DTW 对齐和误差检测：'
    '左/右肩、左/右肘、左/右腕、左/右髋、左/右膝、左/右踝。', indent=True)

add_para(
    '当某一帧未检测到人体时，系统沿用上一帧的骨架数据，保证序列的连续性。', indent=True)

# 3.2 平滑降噪
doc.add_heading('3.2  平滑降噪', level=2)

add_para(
    'MediaPipe 是逐帧独立推理的，即使人体静止不动，相邻帧的骨架坐标也会有微小的随机偏移'
    '（通常 1~3 像素）。如果不处理，会导致骨架线条闪烁、误差检测产生大量误报。', indent=True)

add_para(
    '系统采用滑动平均滤波进行降噪，窗口大小为 5 帧。'
    '对每个关节的每个坐标维度（x、y、z）分别做卷积平均：', indent=True)

add_para('    smoothed[i] = (x[i-2] + x[i-1] + x[i] + x[i+1] + x[i+2]) / 5')

add_para(
    '该方法在消除高频抖动的同时保留了真实的大幅度动作变化，'
    '对老师和学生视频都做同样的处理。', indent=True)

# 3.3 空间归一化
doc.add_heading('3.3  空间归一化', level=2)

add_para(
    '老师和学生的身高、臂长、腿长各不相同，直接用像素坐标比对会产生大量误判。'
    '空间归一化的目标是将骨架转化为一种与体型无关的表示。', indent=True)

add_para('具体做法分为三步：', indent=True)
add_para('    第一步：计算骨盆中心 Root = (左髋坐标 + 右髋坐标) / 2')
add_para('    第二步：计算躯干长度 Scale = ||肩中点 - Root||')
add_para('    第三步：归一化每个关节 normalized = (joint - Root) / Scale')

add_para(
    '归一化后，所有坐标变成"几个躯干长度"的倍数。'
    '无论高矮胖瘦，做出同一个动作后归一化坐标基本一致，从而可以进行有意义的比对。', indent=True)

# 3.4 DTW 时间对齐
doc.add_heading('3.4  DTW 时间对齐', level=2)

add_para(
    '老师和学生做同一套动作时，节奏不可能完全一致——某个动作可能做得更快或更慢。'
    '如果直接逐帧比对，同一时刻两人可能在做不同的动作阶段，导致全是误报。', indent=True)

add_para(
    '系统使用 DTW（Dynamic Time Warping，动态时间规整）算法解决这一问题。'
    'DTW 能找到两个时间序列之间的最优对齐路径，允许一对一、多对一、一对多的帧映射关系。', indent=True)

add_para('具体实现：', indent=True)
add_para(
    '    - 每帧用 12 个核心关节的 (x, y, z) 拼成 36 维特征向量\n'
    '    - 使用 fastdtw 库计算学生序列和老师序列之间的最优对齐路径\n'
    '    - 距离度量采用欧氏距离\n'
    '    - 对齐结果：学生第 i 帧对应老师第 j 帧（j 不一定等于 i）')

add_para(
    '对齐完成后，无论快做慢做，同一动作阶段的帧都能正确匹配到一起。', indent=True)

# 3.5 误差检测
doc.add_heading('3.5  误差检测', level=2)

add_para(
    'DTW 对齐后，系统遍历每一帧的每个核心关节，'
    '计算学生与老师在归一化空间下的欧氏距离。'
    '当距离超过阈值 0.12（即躯干长度的 12%）时，该关节被标记为错误。', indent=True)

add_para(
    '误差检测只对 8 个末端和中间关节生效：手腕、手肘、膝盖、脚踝。'
    '肩和髋作为参考锚点，不参与报错。', indent=True)

doc.add_heading('3.5.1  防抖机制', level=3)

add_para(
    '由于骨架检测存在微小波动，某个关节可能在阈值附近反复触发，'
    '导致视觉标注不停闪烁。系统采用状态保持机制解决这一问题：', indent=True)

add_para(
    '    - 一旦某关节触发错误，错误状态至少保持 10 帧\n'
    '    - 即使后续帧误差变小，也继续显示错误标注\n'
    '    - 10 帧后如果没有新的错误触发，才自动消除')

add_para('该机制有效避免了视觉闪烁，使错误标注稳定可读。', indent=True)

# ==================== 4. 视觉渲染 ====================
doc.add_heading('4  视觉渲染与输出', level=1)

add_para(
    '系统将老师画面和学生画面左右拼接，生成一段对比纠错视频。'
    '输出视频宽度为两段视频宽度之和，高度统一为 600px。', indent=True)

doc.add_heading('4.1  左侧：标准演示', level=2)
add_para(
    '左侧显示老师的原始视频帧，并叠加绿色骨架连线，作为标准动作参考。', indent=True)

doc.add_heading('4.2  右侧：纠错反馈', level=2)
add_para('右侧在学生画面上叠加多个渲染层：', indent=True)

add_para(
    '    （1）老师残影：将老师的骨架经过空间对齐后，以半透明浅绿色（40% 透明度）'
    '绘制在学生画面上，作为"正确动作应该在哪"的参考。')

add_para(
    '    （2）学生骨架：正确部位画橙色连线，错误部位画红色粗线。')

add_para(
    '    （3）纠错箭头：从学生的错误关节位置，画红色箭头指向对齐后的老师正确位置，'
    '直观引导学生调整方向。箭头长度小于 15 像素时不绘制，避免微小偏差导致箭头抖动。')

doc.add_heading('4.3  空间对齐算法', level=2)
add_para(
    '为了将老师骨架"附体"到学生身上，需要进行画面级的空间对齐：', indent=True)

add_para(
    '    1. 计算学生在画面中的骨盆像素位置\n'
    '    2. 计算老师在画面中的骨盆像素位置\n'
    '    3. 计算两人躯干像素长度的比值 scale\n'
    '    4. 对老师的每个关节点：aligned = 学生骨盆 + (老师关节 - 老师骨盆) * scale')

add_para(
    '对齐后，老师骨架的位置和大小与学生匹配，残影效果直观准确。', indent=True)

# ==================== 5. 关键参数 ====================
doc.add_heading('5  关键参数配置', level=1)

add_para('系统中定义了以下关键参数，均在代码顶部配置区集中管理：', indent=True)

table = doc.add_table(rows=5, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['参数名', '默认值', '说明']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

params = [
    ('NORM_DIST_THRESHOLD', '0.12', '归一化空间距离阈值，以躯干长度的 12% 作为容差'),
    ('SMOOTH_WINDOW', '5', '滑动平均窗口大小，消除骨架提取的逐帧抖动'),
    ('HOLD_FRAMES', '10', '错误状态保持帧数，防止标注闪烁'),
    ('TARGET_HEIGHT', '600px', '输出视频高度，等比缩放以兼顾质量与性能'),
]
for row_idx, (name, val, desc) in enumerate(params, 1):
    table.rows[row_idx].cells[0].text = name
    table.rows[row_idx].cells[1].text = val
    table.rows[row_idx].cells[2].text = desc

doc.add_paragraph()  # spacing

# ==================== 6. 技术依赖 ====================
doc.add_heading('6  技术依赖', level=1)

techs = [
    ('MediaPipe', 'Google 开源的人体姿态估计框架，支持 33 个 3D 关键点的实时检测。'),
    ('OpenCV', '视频读写、图像处理、骨架绘制与画面合成。'),
    ('fastdtw', '动态时间规整算法的高效近似实现，用于帧序列对齐。'),
    ('NumPy', '矩阵运算、向量计算、归一化与平滑处理。'),
    ('SciPy', '欧氏距离计算，用于空间误差度量。'),
    ('Pillow (PIL)', '中文字体渲染与半透明文字叠加。'),
]
for name, desc in techs:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}：')
    run.bold = True
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2 = p.add_run(desc)
    run2.font.name = '宋体'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(4)

# ==================== 7. 总结 ====================
doc.add_heading('7  总结', level=1)

add_para(
    '本项目实现了一个端到端的动作姿态纠错系统，从两段原始视频出发，'
    '经过骨架提取、降噪、归一化、时间对齐、误差检测和视觉渲染等步骤，'
    '最终输出一段左右对比的纠错视频。', indent=True)

add_para('系统具有以下特点：', indent=True)

add_para(
    '    （1）全自动：无需人工标注，输入视频即可自动产出纠错结果。\n'
    '    （2）抗干扰：通过空间归一化消除体型差异，通过 DTW 消除速度差异，'
    '通过滑动平均和状态保持消除检测噪声。\n'
    '    （3）可视化：残影叠加、颜色区分、箭头引导等多种视觉手段，'
    '让使用者能直观看到错误位置和正确方向。')

add_para(
    '该系统可应用于体育教学、舞蹈训练、康复训练等领域，'
    '为动作学习提供客观、自动化的视觉反馈。', indent=True)

# ==================== 保存 ====================
out = r"e:\leetcode\last\AI_Pose_Correction_Report.docx"
doc.save(out)
print(f"Report saved: {out}")
